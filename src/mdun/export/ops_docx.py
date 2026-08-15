"""Quill Delta ops → DOCX 导出（保留在线文档级样式）。

约定（与前端编辑器一致）：
- 行级属性挂在新行符 insert 上：header / list / indent / align；
- 字符级属性挂在文本 insert 上：bold / italic / underline / strike / color / background；
- {insert: {pageHead: N}} 表示第 N 页开始（可导出分页符）。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_ALIGN = {"": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}


def _hex_rgb(color: str | None) -> RGBColor | None:
    if not color or not isinstance(color, str) or not color.startswith("#"):
        return None
    try:
        return RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
    except ValueError:
        return None


def export_docx_from_ops(ops: list, out: str | Path, page_break_per_page: bool = False, toc: list | None = None) -> Path:
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "DengXian"
    normal.font.size = Pt(11)

    if toc:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

        head = doc.add_paragraph()
        r = head.add_run("目录")
        r.bold = True
        r.font.size = Pt(16)
        for entry in toc:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Pt(440), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.add_run("  " * max(0, int(entry.get("level", 1)) - 1) + str(entry.get("text", "")))
            p.add_run("\t" + str(entry.get("page", "")))
        doc.add_page_break()

    para = doc.add_paragraph()
    line_attrs: dict = {}
    current_page = 0
    page_started = False

    def flush_line():
        nonlocal para
        if not para.runs and not para.text:
            return
        p = para
        # 行级样式
        header = line_attrs.get("header")
        align = _ALIGN.get(str(line_attrs.get("align", "")) or "", WD_ALIGN_PARAGRAPH.LEFT)
        p.alignment = align
        lst = line_attrs.get("list")
        indent = int(line_attrs.get("indent") or 0)
        if header:
            p.style = doc.styles[f"Heading {header}"] if header in (1, 2, 3) else p.style
        elif lst == "ordered":
            p.style = doc.styles["List Number"] if indent <= 1 else doc.styles["List Number 2"]
        elif lst == "bullet":
            p.style = doc.styles["List Bullet"] if indent <= 1 else doc.styles["List Bullet 2"]
        para = doc.add_paragraph()

    for op in ops:
        if not isinstance(op, dict):
            continue
        ins = op.get("insert")
        attrs = op.get("attributes") or {}
        if isinstance(ins, dict):
            if ins.get("mdunPage") is not None:
                new_page = int(ins["mdunPage"])
                if new_page != current_page:
                    current_page = new_page
                    flush_line()
                    if page_break_per_page and page_started:
                        doc.add_page_break()
                    page_started = True
            elif ins.get("mdunTable") is not None:
                flush_line()
                rows = ins["mdunTable"] or []
                if rows:
                    n_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=0, cols=n_cols)
                    table.style = "Table Grid"
                    for r in rows:
                        cells = table.add_row().cells
                        for i in range(n_cols):
                            cells[i].text = r[i] if i < len(r) else ""
            continue
        if not isinstance(ins, str):
            continue
        parts = ins.split("\n")
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                if attrs.get("bold"):
                    run.bold = True
                if attrs.get("italic"):
                    run.italic = True
                if attrs.get("underline"):
                    run.underline = True
                if attrs.get("strike"):
                    run.strike = True
                color = _hex_rgb(attrs.get("color"))
                if color:
                    run.font.color.rgb = color
                bg = _hex_rgb(attrs.get("background"))
                if bg:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), attrs["background"].lstrip("#"))
                    run._element.get_or_add_rPr().append(shd)
            if i < len(parts) - 1:
                # 换行符：该行结束；行级属性在该 newline 的 attrs 上
                line_attrs = dict(attrs)
                flush_line()
    flush_line()
    doc.save(str(out))
    return out
