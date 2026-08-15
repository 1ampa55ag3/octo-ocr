"""导出：DOCX（FR-6.2）。段落样式 + 标题层级 + 原生表格 + 公式。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from mdun.pipeline import PageData


def _ordered_items(page: PageData) -> list[tuple[str, dict]]:
    """按 y 坐标把段落/表格/公式交错排序，保持阅读顺序。"""
    items: list[tuple[str, dict]] = []
    for para in page.paras:
        if not para.text.strip():
            continue
        items.append(("para", {"y": para.box[1], "kind": para.kind, "text": para.text.strip()}))
    for t in page.tables:
        box = t.get("box", [0, 0, 0, 0])
        items.append(("table", {"y": box[1], "rows": t.get("rows", [])}))
    for f in page.formulas:
        box = f.get("box", [0, 0, 0, 0])
        items.append(("formula", {"y": box[1], "latex": f.get("latex", "")}))
    items.sort(key=lambda x: x[1]["y"])
    return items


def export_docx(project, out: str | Path, page_break_per_page: bool = False) -> Path:
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "DengXian"
    style.font.size = Pt(11)

    first = True
    for page in project.pages:
        if page_break_per_page and not first and page.paras:
            doc.add_page_break()
        first = False
        for kind, item in _ordered_items(page):
            if kind == "para":
                if item["kind"] == "heading":
                    doc.add_heading(item["text"], level=2)
                elif item["kind"] == "page-number":
                    continue
                else:
                    doc.add_paragraph(item["text"])
            elif kind == "table":
                rows = item["rows"]
                if not rows:
                    continue
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=n_cols)
                table.style = "Table Grid"
                for r in rows:
                    cells = table.add_row().cells
                    for i in range(n_cols):
                        cells[i].text = r[i] if i < len(r) else ""
            elif kind == "formula":
                p = doc.add_paragraph()
                run = p.add_run(item["latex"])
                run.italic = True
                run.font.size = Pt(10)
                p.paragraph_format.alignment = 1  # 居中
    doc.save(str(out))
    return out


def export_xlsx(project, out: str | Path) -> Path:
    """导出 Excel：每个表格一个工作表。"""
    from openpyxl import Workbook

    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    n = 0
    for page in project.pages:
        for t in page.tables:
            rows = t.get("rows", [])
            if not rows:
                continue
            ws = wb.create_sheet(title=f"表{n + 1}-p{page.index + 1}")
            for r in rows:
                ws.append(r)
            n += 1
    if n == 0:
        ws = wb.create_sheet(title="Sheet1")
    wb.save(str(out))
    return out
